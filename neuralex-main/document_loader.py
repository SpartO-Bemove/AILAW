"""
Модуль для загрузки и обработки дополнительных юридических документов
"""
import logging
import os
import hashlib
from typing import List, Dict, Optional
from pathlib import Path

# Импорты для обработки документов
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF не установлен, PDF файлы не будут обрабатываться")

try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx не установлен, DOCX файлы не будут обрабатываться")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Класс для загрузки документов из папки documents/"""
    
    def __init__(self, documents_path: str = "documents"):
        self.documents_path = Path(documents_path)
        
        # Определяем поддерживаемые форматы на основе доступных библиотек
        self.supported_extensions = {'.txt', '.md'}
        if PYMUPDF_AVAILABLE:
            self.supported_extensions.add('.pdf')
        if PYTHON_DOCX_AVAILABLE:
            self.supported_extensions.add('.docx')
            
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Создаем папки если их нет
        self._create_directories()
    
    def _create_directories(self):
        """Создает необходимые директории"""
        directories = ['laws', 'codes', 'articles', 'court_practice']
        for directory in directories:
            (self.documents_path / directory).mkdir(parents=True, exist_ok=True)
    
    def get_file_hash(self, file_path: Path) -> str:
        """Получает хэш файла для отслеживания изменений"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception as e:
            logger.error(f"Ошибка при получении хэша файла {file_path}: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: Path) -> Optional[str]:
        """Извлекает текст из файла в зависимости от его типа"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.txt' or extension == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif extension == '.pdf':
                if not PYMUPDF_AVAILABLE:
                    logger.warning(f"PyMuPDF не доступен, пропускаем PDF файл: {file_path}")
                    return None
                
                try:
                    doc = fitz.open(file_path)
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                    return text
                except Exception as e:
                    logger.error(f"Ошибка при чтении PDF {file_path}: {e}")
                    return None
            
            elif extension == '.docx':
                if not PYTHON_DOCX_AVAILABLE:
                    logger.warning(f"python-docx не доступен, пропускаем DOCX файл: {file_path}")
                    return None
                
                try:
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except Exception as e:
                    logger.error(f"Ошибка при чтении DOCX {file_path}: {e}")
                    return None
            
            else:
                logger.warning(f"Неподдерживаемый формат файла: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из {file_path}: {e}")
            return None
    
    def load_documents_from_directory(self, directory: Path, category: str) -> List[Document]:
        """Загружает все документы из указанной директории"""
        documents = []
        
        if not directory.exists():
            logger.info(f"Директория {directory} не существует")
            return documents
        
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    # Проверяем размер файла (максимум 10 МБ)
                    if file_path.stat().st_size > 10 * 1024 * 1024:
                        logger.warning(f"Файл {file_path.name} слишком большой (>10MB), пропускаем")
                        continue
                    
                    text = self.extract_text_from_file(file_path)
                    if text and len(text.strip()) > 50:  # Минимум 50 символов
                        # Разбиваем на чанки
                        chunks = self.text_splitter.split_text(text)
                        
                        for i, chunk in enumerate(chunks):
                            doc = Document(
                                page_content=chunk,
                                metadata={
                                    'source': str(file_path),
                                    'filename': file_path.name,
                                    'category': category,
                                    'chunk_id': i,
                                    'file_hash': self.get_file_hash(file_path)
                                }
                            )
                            documents.append(doc)
                        
                        logger.info(f"✅ Загружен документ: {file_path.name} ({len(chunks)} фрагментов)")
                    else:
                        logger.warning(f"Файл {file_path.name} пустой или слишком короткий")
                        
                except Exception as e:
                    logger.error(f"Ошибка при обработке файла {file_path}: {e}")
        
        return documents
    
    def load_all_documents(self) -> List[Document]:
        """Загружает все документы из всех категорий"""
        all_documents = []
        
        categories = {
            'laws': 'Федеральные законы',
            'codes': 'Кодексы РФ', 
            'articles': 'Юридические статьи',
            'court_practice': 'Судебная практика'
        }
        
        for folder, category in categories.items():
            directory = self.documents_path / folder
            documents = self.load_documents_from_directory(directory, category)
            all_documents.extend(documents)
            
            if documents:
                logger.info(f"📚 Категория '{category}': загружено {len(documents)} фрагментов")
        
        logger.info(f"📊 Всего загружено документов: {len(all_documents)} фрагментов")
        return all_documents
    
    def get_documents_stats(self) -> Dict:
        """Возвращает статистику по загруженным документам"""
        stats = {
            'total_files': 0,
            'categories': {},
            'supported_formats': list(self.supported_extensions)
        }
        
        categories = ['laws', 'codes', 'articles', 'court_practice']
        
        for category in categories:
            directory = self.documents_path / category
            if directory.exists():
                files = [f for f in directory.iterdir() 
                        if f.is_file() and f.suffix.lower() in self.supported_extensions]
                stats['categories'][category] = len(files)
                stats['total_files'] += len(files)
        
        return stats